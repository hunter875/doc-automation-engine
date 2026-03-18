import io

from docx import Document

from app.services.word_scanner import scan_word_template


def _build_docx_bytes(build_fn):
    doc = Document()
    build_fn(doc)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_table_statistics_placeholders_remain_scalar_numbers():
    def build(doc):
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Chỉ tiêu"
        table.cell(0, 1).text = "Kết quả"
        table.cell(1, 0).text = "Tổng cháy"
        table.cell(1, 1).text = "{{ stt_02_tong_chay }}"
        doc.add_paragraph("Tổng nổ: {{ stt_08_tong_no }}")

    result = scan_word_template(_build_docx_bytes(build), use_llm=False)
    fields = {field["name"]: field for field in result["schema_definition"]["fields"]}
    rules = {rule["output_field"]: rule for rule in result["aggregation_rules"]["rules"]}

    assert fields["stt_02_tong_chay"]["type"] == "number"
    assert fields["stt_08_tong_no"]["type"] == "number"
    assert rules["stt_02_tong_chay"]["method"] == "SUM"
    assert rules["stt_08_tong_no"]["method"] == "SUM"



def test_explicit_for_loop_creates_array_object_schema():
    def build(doc):
        doc.add_paragraph("{%p for vu in danh_sach_cnch %}")
        doc.add_paragraph("{{ vu.thoi_gian }} - {{ vu.dia_diem }} - {{ vu.ket_qua_xu_ly }}")
        doc.add_paragraph("{%p endfor %}")

    result = scan_word_template(_build_docx_bytes(build), use_llm=False)
    fields = {field["name"]: field for field in result["schema_definition"]["fields"]}
    rules = {rule["output_field"]: rule for rule in result["aggregation_rules"]["rules"]}

    assert fields["danh_sach_cnch"]["type"] == "array"
    item_fields = {field["name"]: field for field in fields["danh_sach_cnch"]["items"]["fields"]}
    assert set(item_fields) >= {"thoi_gian", "dia_diem", "ket_qua_xu_ly"}
    assert rules["danh_sach_cnch"]["method"] == "CONCAT"



def test_metadata_fields_are_not_added_to_aggregation_rules():
    def build(doc):
        doc.add_paragraph("Ngày xuất: {{ ngay_xuat }}")
        doc.add_paragraph("Tháng xuất: {{ thang_xuat }}")
        doc.add_paragraph("Năm xuất: {{ nam_xuat }}")
        doc.add_paragraph("Tổng cháy: {{ tong_so_vu_chay }}")

    result = scan_word_template(_build_docx_bytes(build), use_llm=False)
    fields = {field["name"]: field for field in result["schema_definition"]["fields"]}
    rules = {rule["output_field"]: rule for rule in result["aggregation_rules"]["rules"]}

    assert set(fields) >= {"ngay_xuat", "thang_xuat", "nam_xuat", "tong_so_vu_chay"}
    assert "ngay_xuat" not in rules
    assert "thang_xuat" not in rules
    assert "nam_xuat" not in rules
    assert rules["tong_so_vu_chay"]["method"] == "SUM"
