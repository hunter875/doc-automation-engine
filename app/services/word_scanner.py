"""Word template scanner — extract placeholders and infer schema.

Design rules:
    1. Scalar placeholders ``{{ foo }}`` remain scalar even when placed in tables.
    2. Arrays are inferred ONLY from explicit Jinja loops: ``{% for item in arr %}``.
    3. Metadata/date fields are excluded from Pandas aggregation rules.
    4. Descriptions are deterministic and human-readable (no noisy context scraping).
    5. STRICT TYPING: Prefixes `stt_` and `tong_` are strictly NUMBER.
"""

import io
import logging
import re
from typing import Any

from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

# Regex ĐÃ ĐƯỢC FIX: Bắt cả {{ bien }} LẪN {%p if bien == 0 %}
PLACEHOLDER_RE = re.compile(
    r"\{\{\s*([A-Za-z_][A-Za-z0-9_\.]*)\s*\}\}|"
    r"\{%-?\s*(?:(?:p|tr|tc)\s+)?if\s+([A-Za-z_][A-Za-z0-9_\.]*)(?:\s*(?:==|!=|>|<|>=|<=)\s*.*?)?\s*-?%\}"
)

LOOP_TOKEN_RE = re.compile(
    r"\{%-?\s*(?:(?:p|tr|tc)\s+)?(for\s+(?P<alias>[A-Za-z_][A-Za-z0-9_]*)\s+in\s+(?P<array>[A-Za-z_][A-Za-z0-9_]*)|endfor)\s*-?%\}"
)

# ── Type-inference keyword lists ──────────────────────────────

NUMBER_KEYWORDS = [
    "so_", "tong_", "amount", "total", "price", "quantity",
    "qty", "count", "sl_", "percent", "rate", "cost", "revenue", "fee",
    "value", "age", "year", "month", "day", "area", "weight",
]

ARRAY_KEYWORDS = [
    "danh_sach", "danh_ds", "list_", "_list", "items_", "_items",
    "bang_", "_bang", "table_", "_table", "rows_", "_rows", "records_",
    "chi_tiet", "details_", "_details", "entries_", "hang_muc",
]

METADATA_KEYWORDS = [
    "ngay_", "thang_", "nam_", "tu_ngay", "den_ngay", "tuan_", "ky_",
    "ngaybao", "bao_cao", "xuat_", "ngayxuat", "thangxuat", "namxuat",
]

# ĐÃ FIX: Xóa chữ "co_" (tránh nhầm với co_so), chỉ giữ các tiền tố rành mạch
BOOLEAN_KEYWORDS = [
    "is_", "has_", "da_", "approved", "active", "enabled",
    "verified", "confirmed",
]

_SUBFIELD_TYPE_HINTS = {
    "so_ky_hieu": "string", 
    "bien_so": "string",    
    "id": "string",
    "ma_": "string",
    "ten_": "string",
    "name_": "string",
    "loai_": "string",
    "type_": "string",
    "dia_chi": "string",
    "address": "string",
    "mo_ta": "string",
    "description": "string",
    "ghi_chu": "string",
    "note": "string",
    "ngay_": "string",
    "date_": "string",
    "thoi_gian": "string",
    "time_": "string",
    "so_": "number",
    "tong_": "number",
    "amount_": "number",
    "count_": "number",
    "price_": "number",
    "qty_": "number",
}

# ── Helpers ────────────────────────────────────────────────────

def _infer_type(var_name: str) -> str:
    """Đã dẹp bỏ trò đoán theo context để tránh ảo giác (Hallucination) type."""
    name_lower = var_name.lower()

    if _is_metadata_field(var_name):
        return "string"

    # LUẬT THÉP SỐ 1: Cứ stt_ hoặc tong_ là Number. Bất chấp tất cả.
    if name_lower.startswith("stt_") or name_lower.startswith("tong_"):
        return "number"

    for kw in BOOLEAN_KEYWORDS:
        if name_lower.startswith(kw):
            return "boolean"
    for kw in ARRAY_KEYWORDS:
        if kw in name_lower:
            return "array"
    for kw in NUMBER_KEYWORDS:
        if kw in name_lower:
            return "number"
    
    return "string"

def _infer_subfield_type(col_name: str) -> str:
    col_lower = col_name.lower()
    
    # Bước 1: Ưu tiên check khớp chính xác 100% trước (để bắt ngoại lệ như bien_so)
    if col_lower in _SUBFIELD_TYPE_HINTS:
        return _SUBFIELD_TYPE_HINTS[col_lower]
        
    # Bước 2: Chỉ check xem có BẮT ĐẦU BẰNG tiền tố đó không
    for prefix, t in _SUBFIELD_TYPE_HINTS.items():
        if col_lower.startswith(prefix):
            return t
            
    # Bước 3: Đéo đoán được thì auto về String cho an toàn
    return "string"

def _humanize_name(name: str) -> str:
    return name.replace("_", " ").strip().capitalize()

def _infer_description(var_name: str) -> str:
    return _humanize_name(var_name)

def _is_metadata_field(var_name: str) -> bool:
    name_lower = var_name.lower()
    return any(keyword in name_lower for keyword in METADATA_KEYWORDS)

def _to_snake_case(name: str) -> str:
    s = re.sub(r"[^a-zA-Z0-9_]", "_", name)
    s = re.sub(r"_+", "_", s).strip("_").lower()
    if s and s[0].isdigit():
        s = "field_" + s
    return s or "unnamed_field"

def _extract_loop_arrays(full_text: str) -> dict[str, dict[str, Any]]:
    arrays: dict[str, dict[str, Any]] = {}
    stack: list[dict[str, Any]] = []

    token_re = re.compile(
        r"\{\{\s*[A-Za-z_][A-Za-z0-9_\.]*\s*\}\}|\{%-?\s*(?:(?:p|tr|tc)\s+)?(?:for\s+[A-Za-z_][A-Za-z0-9_]*\s+in\s+[A-Za-z_][A-Za-z0-9_]*|endfor)\s*-?%\}"
    )

    for match in token_re.finditer(full_text):
        token = match.group(0)
        loop_match = LOOP_TOKEN_RE.match(token)
        if loop_match:
            if token.strip().endswith("endfor %}") or "endfor" in token:
                if stack:
                    loop_info = stack.pop()
                    array_name = loop_info["array_name"]
                    arrays.setdefault(array_name, {
                        "item_alias": loop_info["alias"],
                        "fields": {},
                    })
                    arrays[array_name]["fields"].update(loop_info["fields"])
                continue

            alias = loop_match.group("alias")
            array_name = loop_match.group("array")
            if alias and array_name:
                stack.append({
                    "alias": alias,
                    "array_name": array_name,
                    "fields": {},
                })
            continue

        placeholder_match = re.match(r"\{\{\s*([A-Za-z_][A-Za-z0-9_\.]*)\s*\}\}", token)
        if not placeholder_match:
            continue

        raw_name = placeholder_match.group(1)
        if "." not in raw_name or not stack:
            continue

        alias, _, field_name = raw_name.partition(".")
        for loop_info in reversed(stack):
            if loop_info["alias"] != alias:
                continue
            field_key = _to_snake_case(field_name)
            loop_info["fields"][field_key] = {
                "name": field_key,
                "type": _infer_subfield_type(field_key),
                "description": _humanize_name(field_key),
            }
            break

    normalized: dict[str, dict[str, Any]] = {}
    for array_name, info in arrays.items():
        field_list = list(info["fields"].values())
        if not field_list:
            field_list = [{
                "name": "value",
                "type": "string",
                "description": "Value",
            }]
        normalized[_to_snake_case(array_name)] = {
            "item_alias": info["item_alias"],
            "fields": field_list,
        }
    return normalized

def _extract_all_placeholders(full_text: str) -> list[dict[str, Any]]:
    seen: dict[str, dict[str, Any]] = {}
    for match in PLACEHOLDER_RE.finditer(full_text):
        # ĐÃ FIX: Lấy group 1 (biến trong {{}}) hoặc group 2 (biến trong {% if %})
        raw_name = match.group(1) or match.group(2)
        if not raw_name:
            continue
            
        key = raw_name.strip()
        if key not in seen:
            seen[key] = {
                "raw_name": key,
                "snake_name": _to_snake_case(key),
                "is_nested": "." in key,
                "occurrences": 0,
            }
        seen[key]["occurrences"] += 1
    return list(seen.values())

# ── Main entry point ──────────────────────────────────────────

def scan_word_template(file_bytes: bytes, use_llm: bool = True) -> dict[str, Any]:
    doc = DocxDocument(io.BytesIO(file_bytes))

    all_text_parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            all_text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                all_text_parts.append(" | ".join(row_texts))

    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if hf:
                for para in hf.paragraphs:
                    if para.text.strip():
                        all_text_parts.append(para.text)

    full_text = "\n".join(all_text_parts)

    loop_arrays = _extract_loop_arrays(full_text)
    all_placeholders = _extract_all_placeholders(full_text)

    found: dict[str, dict[str, Any]] = {}
    for match in PLACEHOLDER_RE.finditer(full_text):
        # ĐÃ FIX: Lấy group chuẩn xác
        raw_name = match.group(1) or match.group(2)
        if not raw_name:
            continue
            
        if "." in raw_name:
            continue
        snake_name = _to_snake_case(raw_name)

        if snake_name in found:
            found[snake_name]["occurrences"] += 1
            continue

        field_type = _infer_type(snake_name)

        found[snake_name] = {
            "name": snake_name,
            "original_name": raw_name,
            "type": field_type,
            "description": _infer_description(raw_name),
            "context_snippet": "",
            "occurrences": 1,
        }

    for array_name, info in loop_arrays.items():
        if array_name in found:
            found[array_name]["type"] = "array"
            continue
        found[array_name] = {
            "name": array_name,
            "original_name": array_name,
            "type": "array",
            "description": _humanize_name(array_name),
            "context_snippet": "",
            "occurrences": 1,
        }

    variables = list(found.values())

    fields: list[dict[str, Any]] = []
    for var in variables:
        field: dict[str, Any] = {
            "name": var["name"],
            "type": var["type"],
            "description": var["description"],
        }

        if var["type"] == "array":
            loop_info = loop_arrays.get(var["name"])
            if loop_info and loop_info["fields"]:
                cols = loop_info["fields"]
                field["items"] = {
                    "type": "object",
                    "description": f"Một hàng trong bảng {var['name']}",
                    "fields": cols,
                }
            else:
                field["items"] = {"type": "string", "description": "Phần tử"}

        fields.append(field)

    schema_definition = {"fields": fields}

    agg_rules: list[dict[str, Any]] = []
    for var in variables:
        if _is_metadata_field(var["name"]):
            continue
        label_base = var["description"].split("—")[0].strip()
        
        # LUẬT THÉP SỐ 2: NUMBER -> SUM, ARRAY -> CONCAT, STRING -> LAST
        if var["type"] == "number":
            agg_rules.append({
                "output_field": var["name"],
                "source_field": var["name"],
                "method": "SUM",
                "label": f"Tổng {label_base}",
            })
        elif var["type"] == "array":
            agg_rules.append({
                "output_field": var["name"],
                "source_field": var["name"],
                "method": "CONCAT",
                "label": f"Tổng hợp {label_base}",
            })
        elif var["type"] == "string" or var["type"] == "boolean":
            agg_rules.append({
                "output_field": var["name"],
                "source_field": var["name"],
                "method": "LAST",
                "label": label_base,
            })

    aggregation_rules = {"rules": agg_rules} if agg_rules else None

    return {
        "field_count": len(variables),
        "all_placeholders": all_placeholders,
        "variables": variables,
        "schema_definition": schema_definition,
        "aggregation_rules": aggregation_rules,
        "stats": {
            "total_holes": len(all_placeholders),
            "total_placeholders": sum(v["occurrences"] for v in variables),
            "unique_variables": len(variables),
            "array_with_object_schema": sum(1 for v in variables if v["type"] == "array" and v["name"] in loop_arrays),
            "paragraphs_scanned": len(doc.paragraphs),
            "tables_scanned": len(doc.tables),
            "metadata_fields_excluded_from_aggregation": sum(1 for v in variables if _is_metadata_field(v["name"])),
        },
    }